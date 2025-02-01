import re
import os
import sys
import shutil
import gc
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document as DocxDocument
from langchain_community.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
import chromadb
import pandas as pd
import traceback
from rag_fusion.constants import SOURCE_DIRECTORY, EMBEDDINGS, CHROMA_SETTINGS, PERSIST_DIRECTORY, release_model_memory, DB
from logger_config import configure_logging

logger = configure_logging('logs', 'rag_ingest.log')

def validate_folder_path(folder_path, queue=None):
    try:
        if not os.path.exists(folder_path):
            raise FileNotFoundError(f"The folder path '{folder_path}' does not exist.")
        if not os.access(folder_path, os.R_OK):
            raise PermissionError(f"The folder path '{folder_path}' is not readable.")
    except Exception as e:
        logger.error(f"Error occurred in validate_folder_path: {e}")
        if queue:
            queue.put(f"Error occurred in validate_folder_path: {e}")
        raise

def get_file_paths_from_folder(folder_path, queue=None, extensions=[".pdf", ".txt", ".docx", ".xlsx", ".xls"]):
    try:
        validate_folder_path(folder_path, queue)
        file_paths = [
            os.path.join(root, file)
            for root, _, files in os.walk(folder_path)
            for file in files if any(file.endswith(ext) for ext in extensions)
        ]
        logger.info(f"File Paths Fetched: {len(file_paths)} files found.")
        return file_paths
    except Exception as e:
        logger.error(f"Error occurred while extracting file paths: {e}")
        if queue:
            queue.put(f"Error occurred while extracting file paths: {e}")
        raise


def clean_text(text):
    """Clean text by removing excessive whitespace, special characters, etc."""
    try:
        text = re.sub(r'\s+', ' ', text)  # Remove excessive whitespace
        text = re.sub(r'\.{2,}', '.', text)  # Remove multiple dots
        return text.strip()
    except Exception as e:
        logger.error(f"Error occurred while cleaning text: {e}")
        raise


def load_docx(path):
    """Load text from a .docx file without headers and footers."""
    try:
        doc = DocxDocument(path)
        full_text = [
            para.text for para in doc.paragraphs
            if not para._element.getparent().tag.endswith(('hdr', 'ftr'))
        ]
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error occurred while loading DOCX file '{path}': {e}")
        raise

def load_excel(path):
    """Load text from an Excel file."""
    try:
        # Load Excel file using pandas
        df = pd.read_excel(path, sheet_name=None)  # Load all sheets as a dictionary of DataFrames
        full_text = []

        for sheet_name, sheet_data in df.items():
            # Combine all text content from the sheet
            sheet_text = sheet_data.astype(str).apply(" ".join, axis=1).str.cat(sep=" ")
            full_text.append(f"Sheet: {sheet_name}\n{sheet_text}")

        return "\n".join(full_text)  # Combine all sheets into one string
    except Exception as e:
        logger.error(f"Error occurred while loading Excel file '{path}': {e}")
        raise

def load_documents(file_paths, queue=None):
    """Load and clean documents from files."""
    documents = []
    for path in file_paths:
        try:
            if path.endswith(".pdf"):
                loader = PyPDFLoader(path)
                raw_documents = loader.load()
            elif path.endswith(".txt"):
                loader = TextLoader(path)
                raw_documents = loader.load()
            elif path.endswith(".docx"):
                content = load_docx(path)
                raw_documents = [Document(page_content=content)]
            elif path.endswith((".xlsx", ".xls")):  # Add support for Excel files
                content = load_excel(path)
                raw_documents = [Document(page_content=content)]
            else:
                logger.warning(f"Unsupported file format for: {path}")
                continue

            cleaned_documents = [
                Document(
                    page_content=clean_text(doc.page_content),
                    metadata=doc.metadata
                ) for doc in raw_documents
            ]
            documents.extend(cleaned_documents)
            logger.info(f"Loaded Document: {path}")
        except Exception as e:
            logger.error(f"Error occurred while loading document '{path}': {e}")
            if queue:
                queue.put(f"Error occurred while loading document '{path}': {e}")
    return documents


def split_and_track(documents, queue=None, chunk_size=1000, chunk_overlap=300):
    """Split text into chunks and track metadata."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
        chunks = text_splitter.split_documents(documents)

        for i, chunk in enumerate(chunks):
            try:
                source = chunk.metadata.get("source", "unknown")
                chunk.metadata = {
                    "index": i,
                    "source": source
                }
            except Exception as e:
                logger.error(f"Error occurred while assigning metadata to chunk {i}: {e}")
        return chunks
    except Exception as e:
        logger.error(f"Error occurred while splitting and tracking documents: {e}")
        if queue:
            queue.put(f"Error occurred while splitting and tracking documents. Error: {e}")
        raise


def initialize_chroma_db(split_docs, embedding_model, persist_directory, chroma_settings, queue=None, batch_size=75):
    """Initialize the ChromaDB vector store."""
    global DB
    try:
        chromadb.api.client.SharedSystemClient.clear_system_cache()
        if DB is None:
            raise ValueError("DB is not initialized. Call `initialize_constants` first.")
        
        for i in range(0, len(split_docs), batch_size):
            batch = split_docs[i:i + batch_size]
            DB = DB.from_documents(
                batch,
                embedding_model,
                persist_directory=persist_directory,
                client_settings=chroma_settings
            )
        logger.info("Initialized ChromaDB vector store in batches.")
    except Exception as e:
        logger.error(f"Error occurred while initializing ChromaDB: {e}")
        if queue:
            queue.put(f"Error occurred while initializing ChromaDB. Error: {e}")
        raise


def clear_chroma_db(queue=0):
    """Clear existing embeddings from the vector store."""
    try:
        if os.path.exists(PERSIST_DIRECTORY):
            shutil.rmtree(PERSIST_DIRECTORY)
            logger.info(f"Successfully removed the ChromaDB directory: {PERSIST_DIRECTORY}")
        else:
            logger.info("No existing ChromaDB found to clear.")

        chromadb.api.client.SharedSystemClient.clear_system_cache()
        logger.info("Cleared the DB cache.")
    except Exception as e:
        logger.error(f"Error occurred while clearing ChromaDB: {e}")
        if queue:
            queue.put(f"Error cleaning chroma DB. Error: {e}")
        raise

def process_file(file_paths, queue=None):
    try:
        documents = load_documents(file_paths, queue)
        split_docs = split_and_track(documents, queue)
        initialize_chroma_db(split_docs, EMBEDDINGS, PERSIST_DIRECTORY, CHROMA_SETTINGS, queue, batch_size=75)
    except Exception as e:
        error_message = f"{type(e).__name__}: {str(e)}\n{traceback.format_exc()}"
        # Log and send the error to the queue
        logger.error(f"Error processing file '{file_paths}': {error_message}")
        if queue:
            queue.put(error_message)
        raise

def main(queue=None):
    try:
        clear_chroma_db(queue)
        folder_path = SOURCE_DIRECTORY
        file_paths = get_file_paths_from_folder(folder_path, queue)
        process_file(file_paths, queue)
        release_model_memory()
        logger.info("Processing completed successfully.")
    except Exception as e:
        logger.error(f"Error occurred in main process: {e}")
        if queue:
            queue.put(f"Error occurred in main process: {e}")
        raise


if __name__ == "__main__":
    main()
